# Evenetually need level 3/4 score sheets too


from docx import Document
import json
import dbaccess
from dbaccess import cwagsDBSelect

file = Document('CWAGSScoreSheet.docx')

run_info_for_event = cwagsDBSelect("Select * from running_order where event = 5")

m = 1
runs = run_info_for_event


#get the tables in the file
prop = file.core_properties
sections = file.sections

paragraphs = file.paragraphs
tables = file.tables




tables_with_data = [0]


docxdict = {}
def create_dict_from_db(tables):
    for table_number in tables_with_data:
        current_table = tables[table_number]
        #print current_table.columns
        for row in range(0,len(current_table.rows)):
            docxdict[current_table.row_cells(row)[0].text] = current_table.row_cells(row)[1].text
            #print docxdict


list_of_lists = []
list_of_dicts = []

def create_list_of_lists(runs):
    for run in runs:
        #print run.keys()
        new_list = []
        docxdict = {'Judge':'','CWAGS Number':'','Dog Name':'', 'Handler Name':'', 'Round':'', 'Level':''}
        dog = run['dog']
    
        #print dog
        personnamefind = cwagsDBSelect("Select person.name from person left join dog on dog.owner = person.id where dog.id = :id;", {"id":dog})
        personname = personnamefind.next()['name']
        cwagsnumfind = cwagsDBSelect("Select cwags from dog where id = :id;", {"id":dog})
        cwagsnum = cwagsnumfind.next()['cwags']
        name = run['name']
        level = run['level']
        round = run['round']
    #this probably doesnt work but i dont care
        judgefind = cwagsDBSelect("Select name from person left join judge on judge.id = person.id left join round on round.judge = judge.id;",{'round':round})
         
        judge = judgefind.next()['name']
        new_list = [judge, cwagsnum, name, personname, round, level]
        #print new_list
        list_of_lists.append(new_list)
        print run['id'], run['name'], run['round']
    

def edit_docx_from_list_list(list_of_lists,output_files):
    n=1
    filename = output_files
    fileext = ".docx"
    for dict in list_of_lists:
        current_table = tables[0]    
        #print len(current_table.rows)
        for row in range(1,len(current_table.rows)):
            key = row-1
            current_table.row_cells(row)[1].text = str(dict[key])
            #print dict[key]
        file.save(filename+str(n)+fileext)
        n+=1

create_dict_from_db(tables)
create_list_of_lists(runs)
edit_docx_from_list_list(list_of_lists,"ScribeOct23")
